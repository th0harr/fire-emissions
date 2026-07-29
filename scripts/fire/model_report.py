from __future__ import annotations

"""
Read and report the current Fire Emissions model results.

This script is deliberately separate from ``scripts.model``:

    python -m scripts.model ...
        Builds and writes the model tables.

    python -m scripts.fire.model_report ...
        Reads the existing tables and creates report-ready summaries.

The reporting script never rebuilds or overwrites the model.  It opens the
SQLite database in read-only mode, asks SQLite to perform the aggregation, and
only loads the much smaller summary tables into Python memory.

Current CLI examples
--------------------
Default final-emissions report (Stage 2 + replacement):

    python -m scripts.fire.model_report --profile tom --db fire_db

Group the final outputs by dwelling and fire category:

    python -m scripts.fire.model_report --profile tom --db fire_db \
        --group-by dwelling fire_cat

Create plots for every requested stage:

    python -m scripts.fire.model_report --profile tom --db fire_db --plots

Use the stage-specific default units (also used when --units is omitted):

    python -m scripts.fire.model_report --profile tom --db fire_db --units default

Choose one or both explicit mass units for every summary:

    python -m scripts.fire.model_report --profile tom --db fire_db --units kg tonnes

Write tables and plots to one Word document:

    python -m scripts.fire.model_report --profile tom --db fire_db --write

Notes
-----
* ``--hide-ranges`` is a display option only.  It selects the existing
  ``default`` rows and hides the existing ``low`` and ``high`` rows.  It does
  not rerun the model or change uncertainty propagation.
* Plotting requires matplotlib.
* Word-document output requires python-docx.
"""

import argparse
import math
import sqlite3
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Optional, Sequence

from scripts.path_config import load_local_paths_config, resolve_db_path


# =============================================================================
# CONTROLLED REPORTING VOCABULARY
# =============================================================================

# Logical reporting stages.  These names are user-facing CLI values rather than
# physical database table names.
STAGE_ORDER = [
    "inventory",
    "event",
    "stage1",
    "replacement",
    "stage2",
]

DEFAULT_STAGES = ["stage2", "replacement"]

STAGE_LABELS = {
    "inventory": "Inventory snapshot",
    "event": "Fire events and coverage",
    "stage1": "Stage 1 affected carbon stock",
    "replacement": "Embodied CO2 emissions",
    "stage2": "Stage 2 direct combustion emissions",
}

# Mass units exposed by the reporting CLI.  The model tables retain their
# original base values in kilograms; conversion to tonnes happens only in this
# read-only reporting layer.
EXPLICIT_UNIT_ORDER = ["kg", "tonnes"]
UNIT_ORDER = ["default", *EXPLICIT_UNIT_ORDER]
DEFAULT_UNITS = ["default"]


# Each grouping alias has:
#   * one stable CLI name;
#   * one display label;
#   * the physical column used by each logical stage.
#
# A value of None means that the grouping is not available for that stage.
# ``__virtual__`` is used where the reporting query creates a value rather than
# reading a physical source column, for example Stage 1 carbon_origin.
GROUPING_REGISTRY: dict[str, dict[str, Any]] = {
    "dwelling": {
        "label": "Dwelling type",
        "columns": {
            "inventory": "dwelling_type",
            "event": "dwelling_type_for_model",
            "stage1": "dwelling_type_for_model",
            "replacement": "dwelling_type_for_model",
            "stage2": "dwelling_type_for_model",
        },
    },
    "room": {
        "label": "Room type",
        "columns": {
            "inventory": "room_type",
            "event": "room_of_origin",
            "stage1": "room_of_origin",
            "replacement": "room_of_origin",
            "stage2": "room_of_origin",
        },
    },
    "year": {
        "label": "Fiscal year start",
        "columns": {
            "inventory": None,
            "event": "fiscal_year_start",
            "stage1": "fiscal_year_start",
            "replacement": "fiscal_year_start",
            "stage2": "fiscal_year_start",
        },
    },
    "fire_cat": {
        "label": "Fire spread category",
        "columns": {
            "inventory": None,
            "event": "fire_spread_category",
            "stage1": "fire_spread_category",
            "replacement": "fire_spread_category",
            "stage2": "fire_spread_category",
        },
    },
    "component": {
        "label": "Model component",
        "columns": {
            "inventory": None,
            "event": None,
            "stage1": "component_type",
            "replacement": "component_type",
            "stage2": "component_type",
        },
    },
    "occupancy": {
        "label": "Occupancy",
        "columns": {
            "inventory": None,
            "event": "occupancy",
            "stage1": "occupancy",
            "replacement": "occupancy",
            "stage2": "occupancy",
        },
    },
    "pathway": {
        "label": "Emission pathway",
        "columns": {
            "inventory": None,
            "event": None,
            "stage1": "emission_pathway",
            "replacement": "emission_pathway",
            "stage2": None,
        },
    },
    "species": {
        "label": "Emission species",
        "columns": {
            "inventory": None,
            "event": None,
            "stage1": None,
            "replacement": None,
            "stage2": "emission_species",
        },
    },
    "carbon_origin": {
        "label": "Carbon origin",
        "columns": {
            "inventory": "__virtual__",
            "event": None,
            "stage1": "__virtual__",
            "replacement": None,
            "stage2": "carbon_origin",
        },
    },
    "property": {
        "label": "FRIS property type",
        "columns": {
            "inventory": None,
            "event": "property_type_3_input",
            "stage1": "property_type_3_input",
            "replacement": "property_type_3_input",
            "stage2": "property_type_3_input",
        },
    },
    "item": {
        "label": "Combusted item",
        "columns": {
            "inventory": None,
            "event": "item_combusted",
            "stage1": "item_combusted",
            "replacement": "item_combusted",
            "stage2": None,
        },
    },
}


# Physical table names used by the current model.
TABLE_METADATA = "fire_model_metadata"
TABLE_FIRE_EVENTS = "fire_events"
TABLE_EVENT_OMISSIONS = "fire_event_omission_summary"
TABLE_MODEL_OMISSIONS = "fire_model_omission_summary"
TABLE_INVENTORY_SNAPSHOT = "inventory_snapshot"
TABLE_INVENTORY_ROOMS = "inventory_room_snapshot"
TABLE_INVENTORY_DWELLINGS = "inventory_dwelling_size_snapshot"
TABLE_STAGE1 = "fire_model_stage1_component_results"
TABLE_STAGE2 = "fire_model_stage2_species_results"


# Keep terminal and Word outputs readable when a very detailed grouping is used.
# The full summary remains available in memory; only the displayed/written copy
# is limited.  A message is printed whenever a table is shortened.
MAX_CONSOLE_ROWS = 100
MAX_DOCUMENT_ROWS = 1000
MAX_PLOT_CATEGORIES = 40


# =============================================================================
# SMALL DATA OBJECTS
# =============================================================================

@dataclass(frozen=True)
class AppliedGrouping:
    """One grouping alias resolved to one physical source column."""

    alias: str
    label: str
    column: str


@dataclass
class ReportTable:
    """One report-ready summary table produced by Stage 3."""

    stage: str
    name: str
    title: str
    rows: list[dict[str, Any]]
    columns: list[str]

    # Groupings that were actually used by this particular table.
    applied_groupings: list[str] = field(default_factory=list)

    # Requested groupings that could not be applied to this table.
    skipped_groupings: list[str] = field(default_factory=list)

    # Optional short lines shown above the grouping information.  These are
    # useful for one-off table context such as the number of modelled room or
    # dwelling types, without repeating the same count in every table row.
    summary_lines: list[str] = field(default_factory=list)

    # Optional stage-level incident count shown once above the table rather
    # than repeated on every grouped row.  For Stage 2 this is derived from
    # Stage 1 direct-pathway rows, so zero-valued lower cases are still counted.
    incident_count: Optional[int] = None
    incident_count_label: str = "Contributing incidents"

    # Plot instructions.  Range-bearing tables plot the estimate and, when
    # visible, use the lower/upper columns as error-bar limits.
    plot_value_column: Optional[str] = None
    plot_lower_column: Optional[str] = None
    plot_upper_column: Optional[str] = None
    plot_value_label: Optional[str] = None
    plot_category_columns: list[str] = field(default_factory=list)
    plot_series_column: Optional[str] = None

    # Optional short explanation displayed beneath the table title and above
    # the grouping/count information.  Keep this concise because it appears in
    # both the terminal and the written report document.
    description: Optional[str] = None

    # Optional terminal-width overrides for selected columns.  Most columns
    # retain the standard maximum width, but text-heavy fields such as
    # omission reasons can be allowed more room before truncation.
    console_column_max_widths: dict[str, int] = field(default_factory=dict)

    # Notes are retained as internal documentation metadata.  They are not
    # printed above the terminal or Word summary tables.
    notes: list[str] = field(default_factory=list)


@dataclass
class ReportResult:
    """Complete result returned by the reporting aggregation layer."""

    db_path: Path
    requested_stages: list[str]
    requested_groupings: list[str]
    hide_ranges: bool
    units: list[str]
    metadata: dict[str, Any]
    tables: list[ReportTable]


# =============================================================================
# GENERIC SQLITE HELPERS
# =============================================================================

def _connect_read_only(db_path: Path) -> sqlite3.Connection:
    """
    Open the database in SQLite read-only mode.

    Using read-only mode provides a useful safety guarantee: this reporting
    script cannot accidentally delete, replace, or update model outputs.
    """
    uri = db_path.resolve().as_uri() + "?mode=ro"
    conn = sqlite3.connect(uri, uri=True)
    conn.row_factory = sqlite3.Row
    return conn


def _quote_ident(name: str) -> str:
    """Safely quote a SQLite table or column identifier."""
    return '"' + name.replace('"', '""') + '"'


def _list_tables(conn: sqlite3.Connection) -> set[str]:
    rows = conn.execute(
        """
        SELECT name
        FROM sqlite_master
        WHERE type IN ('table', 'view');
        """
    ).fetchall()
    return {str(row[0]) for row in rows}


def _table_columns(conn: sqlite3.Connection, table: str) -> set[str]:
    rows = conn.execute(
        f"PRAGMA table_info({_quote_ident(table)});"
    ).fetchall()
    return {str(row[1]) for row in rows}


def _require_table(conn: sqlite3.Connection, table: str) -> None:
    if table not in _list_tables(conn):
        raise RuntimeError(
            f"Required reporting source table is missing: {table}. "
            "Build or update the fire database before running model_report."
        )


def _fetch_dicts(
    conn: sqlite3.Connection,
    sql: str,
    params: Sequence[Any] = (),
) -> list[dict[str, Any]]:
    return [dict(row) for row in conn.execute(sql, params).fetchall()]


def _load_latest_metadata(conn: sqlite3.Connection) -> dict[str, Any]:
    """Load the single latest fire-model metadata row, where available."""
    if TABLE_METADATA not in _list_tables(conn):
        return {}

    row = conn.execute(
        f"""
        SELECT *
        FROM {_quote_ident(TABLE_METADATA)}
        ORDER BY metadata_id DESC
        LIMIT 1;
        """
    ).fetchone()

    return dict(row) if row is not None else {}


def _latest_inventory_snapshot_id(
    conn: sqlite3.Connection,
    metadata: dict[str, Any],
) -> int:
    """
    Resolve the snapshot represented by the current model.

    Prefer the snapshot recorded in fire_model_metadata.  If an older metadata
    schema does not contain that value, fall back to the latest snapshot row.
    """
    snapshot_id = metadata.get("inventory_snapshot_id")
    if snapshot_id is not None:
        return int(snapshot_id)

    _require_table(conn, TABLE_INVENTORY_SNAPSHOT)
    row = conn.execute(
        f"""
        SELECT inventory_snapshot_id
        FROM {_quote_ident(TABLE_INVENTORY_SNAPSHOT)}
        ORDER BY inventory_snapshot_id DESC
        LIMIT 1;
        """
    ).fetchone()

    if row is None:
        raise RuntimeError(
            "No inventory snapshot is available. Run the inventory snapshot "
            "workflow before requesting inventory or fire-model reports."
        )

    return int(row[0])


def _input_type_filter(
    table_columns: set[str],
    metadata: dict[str, Any],
    *,
    table_alias: str = "src",
) -> tuple[list[str], list[Any]]:
    """Build an optional input_type filter from the model metadata."""
    input_type = metadata.get("input_type")
    if input_type is None or "input_type" not in table_columns:
        return [], []

    return [f"{table_alias}.{_quote_ident('input_type')} = ?"], [input_type]


# =============================================================================
# CLI NORMALISATION AND GROUPING RESOLUTION
# =============================================================================

def _normalise_stages(values: Sequence[str]) -> list[str]:
    """Expand ``all``, remove duplicates, and preserve requested order."""
    if "all" in values:
        return list(STAGE_ORDER)

    out: list[str] = []
    for value in values:
        if value not in STAGE_ORDER:
            raise ValueError(f"Unknown report stage: {value}")
        if value not in out:
            out.append(value)

    return out


def _normalise_groupings(values: Sequence[str]) -> list[str]:
    """Validate grouping aliases while preserving left-to-right nesting."""
    if not values or (len(values) == 1 and values[0] == "none"):
        return []

    if "none" in values:
        raise ValueError(
            "--group-by none cannot be combined with other grouping values."
        )

    out: list[str] = []
    for value in values:
        if value not in GROUPING_REGISTRY:
            raise ValueError(f"Unknown grouping alias: {value}")
        if value in out:
            raise ValueError(f"Grouping alias was supplied more than once: {value}")
        out.append(value)

    return out


def _resolve_groupings_for_table(
    conn: sqlite3.Connection,
    *,
    stage: str,
    table: str,
    requested_groupings: Sequence[str],
    virtual_aliases: set[str] | None = None,
) -> tuple[list[AppliedGrouping], list[str]]:
    """
    Resolve CLI grouping aliases against one physical source table.

    A logical stage can use several physical tables.  For example, the
    ``inventory`` stage has a room summary and a dwelling summary.  This helper
    therefore checks compatibility at table level and records any skipped
    aliases explicitly.
    """
    virtual_aliases = virtual_aliases or set()
    columns = _table_columns(conn, table)

    applied: list[AppliedGrouping] = []
    skipped: list[str] = []

    for alias in requested_groupings:
        registry_entry = GROUPING_REGISTRY[alias]
        source_column = registry_entry["columns"].get(stage)

        if source_column == "__virtual__":
            if alias in virtual_aliases:
                applied.append(
                    AppliedGrouping(
                        alias=alias,
                        label=str(registry_entry["label"]),
                        column=source_column,
                    )
                )
            else:
                skipped.append(alias)
            continue

        if source_column is None or source_column not in columns:
            skipped.append(alias)
            continue

        applied.append(
            AppliedGrouping(
                alias=alias,
                label=str(registry_entry["label"]),
                column=str(source_column),
            )
        )

    return applied, skipped


def _physical_groupings(
    applied: Sequence[AppliedGrouping],
    *,
    intrinsic_aliases: set[str] | None = None,
) -> list[AppliedGrouping]:
    """Return groupings that need physical SELECT expressions."""
    intrinsic_aliases = intrinsic_aliases or set()
    return [
        grouping
        for grouping in applied
        if grouping.column != "__virtual__"
        and grouping.alias not in intrinsic_aliases
    ]


def _group_projection(
    groupings: Sequence[AppliedGrouping],
    *,
    table_alias: str = "src",
) -> str:
    """Create the SQL expressions used for requested report groupings.

    Room values need a small amount of display cleaning because unresolved
    room fields can be stored as either SQL NULL or an empty string.  Reporting
    both forms as ``unknown`` avoids a visually blank grouping row while leaving
    the stored model data unchanged.
    """
    expressions: list[str] = []

    for grouping in groupings:
        source_column = f"{table_alias}.{_quote_ident(grouping.column)}"

        if grouping.alias == "room":
            expression = (
                "COALESCE(NULLIF(TRIM(CAST("
                f"{source_column} AS TEXT)), ''), 'unknown')"
            )
        else:
            expression = source_column

        expressions.append(
            f"{expression} AS {_quote_ident(grouping.alias)}"
        )

    return ", ".join(expressions)


def _group_names(groupings: Sequence[AppliedGrouping]) -> list[str]:
    return [grouping.alias for grouping in groupings]


def _join_sql_parts(parts: Iterable[str]) -> str:
    """Join non-empty SQL fragments with commas."""
    return ", ".join(part for part in parts if part)


def _group_by_sql(columns: Sequence[str]) -> str:
    if not columns:
        return ""
    return "GROUP BY " + ", ".join(_quote_ident(column) for column in columns)


def _order_by_sql(columns: Sequence[str]) -> str:
    if not columns:
        return ""
    return "ORDER BY " + ", ".join(_quote_ident(column) for column in columns)



# =============================================================================
# RANGE AND UNIT HELPERS
# =============================================================================

def _normalise_units(values: Sequence[str]) -> list[str]:
    """Validate, deduplicate, and preserve the requested unit mode.

    ``default`` is a complete mode rather than an additional physical unit, so
    it cannot be combined with ``kg`` or ``tonnes``.  Keeping this validation in
    one helper also protects callers that use ``build_model_report()`` directly
    rather than going through the command-line parser.
    """
    units: list[str] = []
    for raw_value in values:
        value = str(raw_value).strip().lower()
        if value not in UNIT_ORDER:
            raise ValueError(
                f"Unknown unit {raw_value!r}. Allowed values: "
                + ", ".join(UNIT_ORDER)
            )
        if value not in units:
            units.append(value)

    units = units or list(DEFAULT_UNITS)
    if "default" in units and len(units) > 1:
        raise ValueError(
            "'default' cannot be combined with 'kg' or 'tonnes'. Use either "
            "'--units default' or explicit unit values."
        )
    return units


def _units_for_summary(
    requested_units: Sequence[str],
    *,
    stage: str,
    summary_name: str,
) -> list[str]:
    """Resolve the physical display units for one summary table.

    Under the ``default`` unit mode, room-level inventory carbon stock remains
    in kgC because its values are comparatively small and are most naturally
    read in kilograms of carbon.  All other mass summaries use tonnes because
    their aggregated values are much larger.

    Explicit ``--units kg``, ``--units tonnes`` or ``--units kg tonnes`` values
    override these stage-specific defaults everywhere.
    """
    normalised = _normalise_units(requested_units)
    if normalised != ["default"]:
        return normalised

    if stage == "inventory" and summary_name == "carbon_stock":
        return ["kg"]

    return ["tonnes"]


def _unit_divisor(unit: str) -> float:
    """Return the divisor used to convert a kilogram source value."""
    if unit == "kg":
        return 1.0
    if unit == "tonnes":
        return 1000.0
    raise ValueError(f"Unsupported mass unit: {unit}")


def _range_column_names(
    *,
    unit_stems: dict[str, str],
    units: Sequence[str],
    hide_ranges: bool,
) -> list[str]:
    """Build estimate/lower/upper output column names in display order."""
    columns: list[str] = []
    for unit in units:
        stem = unit_stems[unit]
        columns.append(f"{stem}_estimate")
        if not hide_ranges:
            columns.extend([f"{stem}_lower", f"{stem}_upper"])
    return columns


def _range_sum_expressions(
    *,
    value_expression: str,
    unit_stems: dict[str, str],
    units: Sequence[str],
    hide_ranges: bool,
    estimate_case_expression: str = "src.estimate_case",
) -> list[str]:
    """
    Return SQL expressions that pivot low/default/high rows into columns.

    The database stores estimate cases as rows.  Reporting is easier to compare
    when the default estimate and its lower/upper limits appear on the same row.
    """
    expressions: list[str] = []
    case_map = [("default", "estimate")]
    if not hide_ranges:
        case_map.extend([("low", "lower"), ("high", "upper")])

    for unit in units:
        divisor = _unit_divisor(unit)
        stem = unit_stems[unit]
        for estimate_case, suffix in case_map:
            expressions.append(
                "COALESCE(SUM(CASE "
                f"WHEN {estimate_case_expression} = '{estimate_case}' "
                f"THEN ({value_expression}) / {divisor} END), 0.0) "
                f"AS {_quote_ident(f'{stem}_{suffix}')}"
            )
    return expressions


def _count_distinct_incidents(
    conn: sqlite3.Connection,
    *,
    table: str,
    where_parts: Sequence[str],
    params: Sequence[Any],
) -> int:
    """Count distinct non-blank incidents for one logical reporting stream."""
    where_sql = " AND ".join(where_parts) if where_parts else "1 = 1"
    row = conn.execute(
        f"""
        SELECT COUNT(DISTINCT src.incident_id)
        FROM {_quote_ident(table)} AS src
        WHERE {where_sql}
          AND src.incident_id IS NOT NULL
          AND TRIM(CAST(src.incident_id AS TEXT)) <> '';
        """,
        list(params),
    ).fetchone()
    return int(row[0] or 0) if row is not None else 0


def _count_distinct_snapshot_values(
    conn: sqlite3.Connection,
    *,
    table: str,
    column: str,
    snapshot_id: int,
) -> int:
    """Count distinct, non-blank values in one inventory snapshot column."""
    row = conn.execute(
        f"""
        SELECT COUNT(DISTINCT src.{_quote_ident(column)})
        FROM {_quote_ident(table)} AS src
        WHERE src.inventory_snapshot_id = ?
          AND src.{_quote_ident(column)} IS NOT NULL
          AND TRIM(CAST(src.{_quote_ident(column)} AS TEXT)) <> '';
        """,
        (snapshot_id,),
    ).fetchone()
    return int(row[0] or 0) if row is not None else 0

# =============================================================================
# PUBLIC SUMMARY BUILD
# =============================================================================

def build_model_report(
    db_path: str | Path,
    *,
    stages: Sequence[str] = DEFAULT_STAGES,
    group_by: Sequence[str] = ("none",),
    hide_ranges: bool = False,
    units: Sequence[str] = DEFAULT_UNITS,
) -> ReportResult:
    """
    Build report-ready summaries from the existing fire-model tables.

    This is the public Stage 3 entry point.  It does not display, plot, or write
    the results; those presentation steps are handled later in this module.
    """
    db_path = Path(db_path)
    requested_stages = _normalise_stages(stages)
    requested_groupings = _normalise_groupings(group_by)
    requested_units = _normalise_units(units)

    conn = _connect_read_only(db_path)
    try:
        metadata = _load_latest_metadata(conn)
        tables: list[ReportTable] = []

        for stage in requested_stages:
            if stage == "inventory":
                tables.extend(
                    _build_inventory_tables(
                        conn,
                        metadata=metadata,
                        requested_groupings=requested_groupings,
                        hide_ranges=hide_ranges,
                        units=requested_units,
                    )
                )
            elif stage == "event":
                tables.extend(
                    _build_event_tables(
                        conn,
                        metadata=metadata,
                        requested_groupings=requested_groupings,
                    )
                )
            elif stage == "stage1":
                tables.append(
                    _build_stage1_table(
                        conn,
                        metadata=metadata,
                        requested_groupings=requested_groupings,
                        hide_ranges=hide_ranges,
                        units=_units_for_summary(
                            requested_units,
                            stage="stage1",
                            summary_name="affected_carbon",
                        ),
                    )
                )
            elif stage == "replacement":
                tables.append(
                    _build_replacement_table(
                        conn,
                        metadata=metadata,
                        requested_groupings=requested_groupings,
                        hide_ranges=hide_ranges,
                        units=_units_for_summary(
                            requested_units,
                            stage="replacement",
                            summary_name="replacement_CO2",
                        ),
                    )
                )
            elif stage == "stage2":
                tables.append(
                    _build_stage2_table(
                        conn,
                        metadata=metadata,
                        requested_groupings=requested_groupings,
                        hide_ranges=hide_ranges,
                        units=_units_for_summary(
                            requested_units,
                            stage="stage2",
                            summary_name="direct_emissions",
                        ),
                    )
                )
            else:  # Defensive guard; CLI validation should prevent this.
                raise ValueError(f"Unsupported report stage: {stage}")

        return ReportResult(
            db_path=db_path,
            requested_stages=requested_stages,
            requested_groupings=requested_groupings,
            hide_ranges=hide_ranges,
            units=requested_units,
            metadata=metadata,
            tables=tables,
        )
    finally:
        conn.close()

# =============================================================================
# INVENTORY SUMMARY BUILDERS
# =============================================================================

def _build_inventory_tables(
    conn: sqlite3.Connection,
    *,
    metadata: dict[str, Any],
    requested_groupings: Sequence[str],
    hide_ranges: bool,
    units: Sequence[str],
) -> list[ReportTable]:
    """Build the current set of inventory reporting summaries."""
    _require_table(conn, TABLE_INVENTORY_ROOMS)
    _require_table(conn, TABLE_INVENTORY_DWELLINGS)

    snapshot_id = _latest_inventory_snapshot_id(conn, metadata)

    carbon_units = _units_for_summary(
        units,
        stage="inventory",
        summary_name="carbon_stock",
    )
    replacement_units = _units_for_summary(
        units,
        stage="inventory",
        summary_name="replacement_stock",
    )

    return [
        _build_inventory_carbon_table(
            conn,
            snapshot_id=snapshot_id,
            requested_groupings=requested_groupings,
            hide_ranges=hide_ranges,
            units=carbon_units,
        ),
        _build_inventory_embodied_table(
            conn,
            snapshot_id=snapshot_id,
            requested_groupings=requested_groupings,
            hide_ranges=hide_ranges,
            units=replacement_units,
        ),
        _build_inventory_dwelling_table(
            conn,
            snapshot_id=snapshot_id,
            requested_groupings=requested_groupings,
        ),
    ]

def _build_inventory_carbon_table(
    conn: sqlite3.Connection,
    *,
    snapshot_id: int,
    requested_groupings: Sequence[str],
    hide_ranges: bool,
    units: Sequence[str],
) -> ReportTable:
    applied, skipped = _resolve_groupings_for_table(
        conn,
        stage="inventory",
        table=TABLE_INVENTORY_ROOMS,
        requested_groupings=requested_groupings,
        virtual_aliases={"carbon_origin"},
    )

    physical = _physical_groupings(
        applied,
        intrinsic_aliases={"carbon_origin"},
    )
    group_names = _group_names(physical)
    group_projection = _group_projection(physical)

    origins = [
        (
            "total",
            "q25_total_carbon_kgC",
            "expected_total_carbon_kgC",
            "q75_total_carbon_kgC",
        ),
        (
            "biogenic",
            "q25_biog_carbon_kgC",
            "expected_biog_carbon_kgC",
            "q75_biog_carbon_kgC",
        ),
        (
            "fossil",
            "q25_fossil_carbon_kgC",
            "expected_fossil_carbon_kgC",
            "q75_fossil_carbon_kgC",
        ),
    ]

    union_parts: list[str] = []
    params: list[Any] = []
    for origin, low_column, estimate_column, upper_column in origins:
        projection = _join_sql_parts([
            group_projection,
            "? AS carbon_origin",
            f"src.{_quote_ident(low_column)} AS lower_kgC",
            f"src.{_quote_ident(estimate_column)} AS estimate_kgC",
            f"src.{_quote_ident(upper_column)} AS upper_kgC",
            "src.room_type AS inventory_row_key",
        ])
        union_parts.append(
            f"""
            SELECT {projection}
            FROM {_quote_ident(TABLE_INVENTORY_ROOMS)} AS src
            WHERE src.inventory_snapshot_id = ?
            """
        )
        params.extend([origin, snapshot_id])

    unit_stems = {
        "kg": "carbon_stock_kgC",
        "tonnes": "carbon_stock_tonnesC",
    }
    measure_sql: list[str] = []
    for unit in units:
        divisor = _unit_divisor(unit)
        stem = unit_stems[unit]
        measure_sql.append(
            f"SUM(COALESCE(estimate_kgC, 0.0)) / {divisor} AS {_quote_ident(stem + '_estimate')}"
        )
        if not hide_ranges:
            measure_sql.extend([
                f"SUM(COALESCE(lower_kgC, 0.0)) / {divisor} AS {_quote_ident(stem + '_lower')}",
                f"SUM(COALESCE(upper_kgC, 0.0)) / {divisor} AS {_quote_ident(stem + '_upper')}",
            ])

    outer_select = _join_sql_parts([
        ", ".join(_quote_ident(name) for name in group_names),
        "carbon_origin",
        *measure_sql,
    ])
    group_columns = [*group_names, "carbon_origin"]

    sql = f"""
        WITH inventory_carbon AS (
            {' UNION ALL '.join(union_parts)}
        )
        SELECT {outer_select}
        FROM inventory_carbon
        {_group_by_sql(group_columns)}
        {_order_by_sql(group_columns)};
    """

    rows = _fetch_dicts(conn, sql, params)
    range_columns = _range_column_names(
        unit_stems=unit_stems,
        units=units,
        hide_ranges=hide_ranges,
    )
    plot_stem = unit_stems[units[0]]
    room_type_count = _count_distinct_snapshot_values(
        conn,
        table=TABLE_INVENTORY_ROOMS,
        column="room_type",
        snapshot_id=snapshot_id,
    )
    return ReportTable(
        stage="inventory",
        name="inventory_carbon_stock",
        title="Inventory room carbon stock",
        rows=rows,
        columns=[
            *group_names,
            "carbon_origin",
            *range_columns,
        ],
        summary_lines=[f"Number of room types: {room_type_count:,}"],
        applied_groupings=_group_names(applied),
        skipped_groupings=skipped,
        plot_value_column=f"{plot_stem}_estimate",
        plot_lower_column=None if hide_ranges else f"{plot_stem}_lower",
        plot_upper_column=None if hide_ranges else f"{plot_stem}_upper",
        plot_value_label=(
            "Carbon stock (kgC)" if units[0] == "kg" else "Carbon stock (tonnes C)"
        ),
        plot_category_columns=[*group_names, "carbon_origin"],
        plot_series_column=None,
        notes=[
            "Low/default/high use the q25/expected/q75 room-stock fields already stored in the inventory snapshot.",
            "Total carbon overlaps with the biogenic and fossil rows; do not sum all three carbon origins together.",
        ],
    )

def _build_inventory_embodied_table(
    conn: sqlite3.Connection,
    *,
    snapshot_id: int,
    requested_groupings: Sequence[str],
    hide_ranges: bool,
    units: Sequence[str],
) -> ReportTable:
    applied, skipped = _resolve_groupings_for_table(
        conn,
        stage="inventory",
        table=TABLE_INVENTORY_ROOMS,
        requested_groupings=requested_groupings,
    )
    physical = _physical_groupings(applied)
    group_names = _group_names(physical)
    group_projection = _group_projection(physical)

    unit_stems = {
        "kg": "replacement_CO2_kg",
        "tonnes": "replacement_CO2_tonnes",
    }
    measure_sql: list[str] = []
    for unit in units:
        divisor = _unit_divisor(unit)
        stem = unit_stems[unit]
        measure_sql.append(
            f"SUM(COALESCE(src.expected_embodied_CO2_kg, 0.0)) / {divisor} AS {_quote_ident(stem + '_estimate')}"
        )
        if not hide_ranges:
            measure_sql.extend([
                f"SUM(COALESCE(src.q25_embodied_CO2_kg, 0.0)) / {divisor} AS {_quote_ident(stem + '_lower')}",
                f"SUM(COALESCE(src.q75_embodied_CO2_kg, 0.0)) / {divisor} AS {_quote_ident(stem + '_upper')}",
            ])

    select_sql = _join_sql_parts([
        group_projection,
        *measure_sql,
    ])

    sql = f"""
        SELECT {select_sql}
        FROM {_quote_ident(TABLE_INVENTORY_ROOMS)} AS src
        WHERE src.inventory_snapshot_id = ?
        {_group_by_sql(group_names)}
        {_order_by_sql(group_names)};
    """

    rows = _fetch_dicts(conn, sql, [snapshot_id])
    range_columns = _range_column_names(
        unit_stems=unit_stems,
        units=units,
        hide_ranges=hide_ranges,
    )
    plot_stem = unit_stems[units[0]]
    room_type_count = _count_distinct_snapshot_values(
        conn,
        table=TABLE_INVENTORY_ROOMS,
        column="room_type",
        snapshot_id=snapshot_id,
    )
    return ReportTable(
        stage="inventory",
        name="inventory_replacement_CO2",
        title="Inventory room embodied CO2 stock",
        rows=rows,
        columns=[
            *group_names,
            *range_columns,
        ],
        summary_lines=[f"Number of room types: {room_type_count:,}"],
        applied_groupings=_group_names(applied),
        skipped_groupings=skipped,
        plot_value_column=f"{plot_stem}_estimate",
        plot_lower_column=None if hide_ranges else f"{plot_stem}_lower",
        plot_upper_column=None if hide_ranges else f"{plot_stem}_upper",
        plot_value_label=(
            "Replacement CO2 (kg)" if units[0] == "kg" else "Replacement CO2 (tonnes)"
        ),
        plot_category_columns=group_names,
        plot_series_column=None,
        description=(
            "These are the estimated CO2 emissions used to produce the items "
            "(spend-based approach), normalised to half of their expected lifecycle."
        ),
        notes=[
            "This is the embodied CO2 represented by the room inventory snapshot, before a fire-specific damage fraction is applied."
        ],
    )

def _build_inventory_dwelling_table(
    conn: sqlite3.Connection,
    *,
    snapshot_id: int,
    requested_groupings: Sequence[str],
) -> ReportTable:
    applied, skipped = _resolve_groupings_for_table(
        conn,
        stage="inventory",
        table=TABLE_INVENTORY_DWELLINGS,
        requested_groupings=requested_groupings,
    )
    physical = _physical_groupings(applied)
    group_names = _group_names(physical)
    group_projection = _group_projection(physical)

    weighted_mean_sql = """
        CASE
            WHEN SUM(COALESCE(src.dwelling_type_pmf, 0.0)) > 0
            THEN SUM(
                COALESCE(src.dwelling_size_m2, 0.0)
                * COALESCE(src.dwelling_type_pmf, 0.0)
            ) / SUM(COALESCE(src.dwelling_type_pmf, 0.0))
            ELSE AVG(src.dwelling_size_m2)
        END AS weighted_mean_dwelling_size_m2
    """

    select_sql = _join_sql_parts([
        group_projection,
        "SUM(COALESCE(src.count_value, 0)) AS survey_count",
        "SUM(COALESCE(src.dwelling_type_pmf, 0.0)) AS dwelling_type_pmf_sum",
        weighted_mean_sql,
    ])

    sql = f"""
        SELECT {select_sql}
        FROM {_quote_ident(TABLE_INVENTORY_DWELLINGS)} AS src
        WHERE src.inventory_snapshot_id = ?
        {_group_by_sql(group_names)}
        {_order_by_sql(group_names)};
    """

    rows = _fetch_dicts(conn, sql, [snapshot_id])
    dwelling_type_count = _count_distinct_snapshot_values(
        conn,
        table=TABLE_INVENTORY_DWELLINGS,
        column="dwelling_type",
        snapshot_id=snapshot_id,
    )
    show_dwelling_type_count = not any(
        grouping.alias == "dwelling" for grouping in applied
    )

    return ReportTable(
        stage="inventory",
        name="inventory_dwelling_summary",
        title="Inventory dwelling-size and survey summary",
        rows=rows,
        columns=[
            *group_names,
            "survey_count",
            "dwelling_type_pmf_sum",
            "weighted_mean_dwelling_size_m2",
        ],
        summary_lines=(
            [f"Number of dwelling types: {dwelling_type_count:,}"]
            if show_dwelling_type_count
            else []
        ),
        applied_groupings=_group_names(applied),
        skipped_groupings=skipped,
        plot_value_column="survey_count",
        plot_value_label="Survey count",
        plot_category_columns=group_names,
        plot_series_column=None,
    )


# =============================================================================
# EVENT SUMMARY BUILDERS
# =============================================================================

def _build_event_tables(
    conn: sqlite3.Connection,
    *,
    metadata: dict[str, Any],
    requested_groupings: Sequence[str],
) -> list[ReportTable]:
    _require_table(conn, TABLE_FIRE_EVENTS)

    tables = [
        _build_event_count_table(
            conn,
            metadata=metadata,
            requested_groupings=requested_groupings,
        )
    ]

    if TABLE_EVENT_OMISSIONS in _list_tables(conn):
        tables.append(_build_upstream_omission_table(conn, metadata=metadata))

    if TABLE_MODEL_OMISSIONS in _list_tables(conn):
        tables.append(_build_model_omission_table(conn, metadata=metadata))

    return tables


def _build_event_count_table(
    conn: sqlite3.Connection,
    *,
    metadata: dict[str, Any],
    requested_groupings: Sequence[str],
) -> ReportTable:
    applied, skipped = _resolve_groupings_for_table(
        conn,
        stage="event",
        table=TABLE_FIRE_EVENTS,
        requested_groupings=requested_groupings,
    )
    physical = _physical_groupings(applied)
    group_names = _group_names(physical)
    group_projection = _group_projection(physical)

    table_columns = _table_columns(conn, TABLE_FIRE_EVENTS)
    where_parts, params = _input_type_filter(
        table_columns,
        metadata,
        table_alias="src",
    )
    where_sql = "WHERE " + " AND ".join(where_parts) if where_parts else ""

    if "omit_from_model" in table_columns:
        modelled_count_sql = """
            SUM(
                CASE
                    WHEN LOWER(COALESCE(src.omit_from_model, 'no')) = 'no'
                    THEN 1 ELSE 0
                END
            ) AS model_facing_event_rows
        """
        marked_omit_sql = """
            SUM(
                CASE
                    WHEN LOWER(COALESCE(src.omit_from_model, 'no')) = 'yes'
                    THEN 1 ELSE 0
                END
            ) AS unresolved_event_rows
        """
    else:
        modelled_count_sql = "COUNT(*) AS model_facing_event_rows"
        marked_omit_sql = "0 AS unresolved_event_rows"

    select_sql = _join_sql_parts([
        group_projection,
        modelled_count_sql,
        marked_omit_sql,
    ])

    sql = f"""
        SELECT {select_sql}
        FROM {_quote_ident(TABLE_FIRE_EVENTS)} AS src
        {where_sql}
        {_group_by_sql(group_names)}
        {_order_by_sql(group_names)};
    """

    rows = _fetch_dicts(conn, sql, params)

    return ReportTable(
        stage="event",
        name="fire_event_summary",
        title="Resolved fire-event summary",
        rows=rows,
        columns=[
            *group_names,
            "model_facing_event_rows",
            "unresolved_event_rows",
        ],
        applied_groupings=_group_names(applied),
        skipped_groupings=skipped,
        plot_value_column="model_facing_event_rows",
        plot_value_label="Model-facing events",
        plot_category_columns=group_names,
        plot_series_column=None,
    )


def _build_upstream_omission_table(
    conn: sqlite3.Connection,
    *,
    metadata: dict[str, Any],
) -> ReportTable:
    columns = _table_columns(conn, TABLE_EVENT_OMISSIONS)
    where_parts, params = _input_type_filter(
        columns,
        metadata,
        table_alias="src",
    )
    where_sql = "WHERE " + " AND ".join(where_parts) if where_parts else ""

    sql = f"""
        SELECT
            src.omit_reason,
            SUM(COALESCE(src.omitted_count, 0)) AS omitted_count
        FROM {_quote_ident(TABLE_EVENT_OMISSIONS)} AS src
        {where_sql}
        GROUP BY src.omit_reason
        ORDER BY omitted_count DESC, src.omit_reason;
    """
    rows = _fetch_dicts(conn, sql, params)

    return ReportTable(
        stage="event",
        name="upstream_omissions",
        title="Upstream fire-event omissions",
        rows=rows,
        columns=["omit_reason", "omitted_count"],
        console_column_max_widths={"omit_reason": 96},
        plot_value_column="omitted_count",
        plot_value_label="Omitted incidents",
        plot_category_columns=["omit_reason"],
        notes=[
            "These incidents were excluded before entering fire_events; they are not treated as zero-emission events."
        ],
    )


def _first_existing(columns: set[str], candidates: Sequence[str]) -> Optional[str]:
    for candidate in candidates:
        if candidate in columns:
            return candidate
    return None


def _build_model_omission_table(
    conn: sqlite3.Connection,
    *,
    metadata: dict[str, Any],
) -> ReportTable:
    """
    Read the protective model-stage omission table adaptively.

    The table was introduced after the first schema version, so the code accepts
    a small set of sensible aliases for its stage/reason/count columns.
    """
    columns = _table_columns(conn, TABLE_MODEL_OMISSIONS)

    stage_column = _first_existing(
        columns,
        ["model_stage", "stage", "omission_stage"],
    )
    reason_column = _first_existing(
        columns,
        ["omit_reason", "omission_reason", "reason"],
    )
    count_column = _first_existing(
        columns,
        ["omitted_count", "omit_count", "row_count"],
    )

    if reason_column is None or count_column is None:
        return ReportTable(
            stage="event",
            name="model_stage_omissions",
            title="Model-stage omissions",
            rows=[],
            columns=["model_stage", "omit_reason", "omitted_count"],
            notes=[
                "fire_model_omission_summary exists, but its reason/count columns could not be recognised by the current reporting aliases."
            ],
        )

    select_stage = (
        f"src.{_quote_ident(stage_column)} AS model_stage"
        if stage_column is not None
        else "'unspecified' AS model_stage"
    )

    where_parts, params = _input_type_filter(
        columns,
        metadata,
        table_alias="src",
    )
    where_sql = "WHERE " + " AND ".join(where_parts) if where_parts else ""

    group_sql = (
        f"GROUP BY src.{_quote_ident(stage_column)}, "
        f"src.{_quote_ident(reason_column)}"
        if stage_column is not None
        else f"GROUP BY src.{_quote_ident(reason_column)}"
    )

    sql = f"""
        SELECT
            {select_stage},
            src.{_quote_ident(reason_column)} AS omit_reason,
            SUM(COALESCE(src.{_quote_ident(count_column)}, 0)) AS omitted_count
        FROM {_quote_ident(TABLE_MODEL_OMISSIONS)} AS src
        {where_sql}
        {group_sql}
        ORDER BY omitted_count DESC, model_stage, omit_reason;
    """

    rows = _fetch_dicts(conn, sql, params)

    return ReportTable(
        stage="event",
        name="model_stage_omissions",
        title="Model-stage omissions",
        rows=rows,
        columns=["model_stage", "omit_reason", "omitted_count"],
        plot_value_column="omitted_count",
        plot_value_label="Omitted incidents",
        plot_category_columns=["model_stage", "omit_reason"],
        notes=[
            "The current dataset may legitimately produce zero rows here; the table is retained as a protective reporting route."
        ],
    )


# =============================================================================
# STAGE 1, REPLACEMENT, AND STAGE 2 BUILDERS
# =============================================================================

def _range_where_clause(
    *,
    hide_ranges: bool,
    table_alias: str = "src",
) -> tuple[list[str], list[Any]]:
    if not hide_ranges:
        return [], []
    return [f"{table_alias}.estimate_case = ?"], ["default"]


def _calculation_status_clause(
    columns: set[str],
    *,
    table_alias: str = "src",
) -> list[str]:
    if "calculation_status" not in columns:
        return []
    return [
        f"LOWER(COALESCE({table_alias}.calculation_status, 'ok')) = 'ok'"
    ]


def _build_stage1_table(
    conn: sqlite3.Connection,
    *,
    metadata: dict[str, Any],
    requested_groupings: Sequence[str],
    hide_ranges: bool,
    units: Sequence[str],
) -> ReportTable:
    _require_table(conn, TABLE_STAGE1)

    applied, skipped = _resolve_groupings_for_table(
        conn,
        stage="stage1",
        table=TABLE_STAGE1,
        requested_groupings=requested_groupings,
        virtual_aliases={"carbon_origin"},
    )
    physical = _physical_groupings(
        applied,
        intrinsic_aliases={"carbon_origin"},
    )
    group_names = _group_names(physical)
    group_projection = _group_projection(physical)

    columns = _table_columns(conn, TABLE_STAGE1)
    where_parts = ["LOWER(COALESCE(src.emission_pathway, '')) = 'direct'"]
    params: list[Any] = []

    input_where, input_params = _input_type_filter(columns, metadata)
    where_parts.extend(input_where)
    where_parts.extend(_calculation_status_clause(columns))
    params.extend(input_params)

    origins = [
        ("total", "direct_total_kgC"),
        ("biogenic", "direct_biogenic_kgC"),
        ("fossil", "direct_fossil_kgC"),
    ]

    union_parts: list[str] = []
    union_params: list[Any] = []
    for carbon_origin, value_column in origins:
        projection = _join_sql_parts([
            group_projection,
            "src.estimate_case",
            "? AS carbon_origin",
            "src.incident_id",
            f"src.{_quote_ident(value_column)} AS affected_kgC",
        ])
        union_parts.append(
            f"""
            SELECT {projection}
            FROM {_quote_ident(TABLE_STAGE1)} AS src
            WHERE {' AND '.join(where_parts)}
            """
        )
        union_params.extend([carbon_origin, *params])

    unit_stems = {
        "kg": "affected_kgC",
        "tonnes": "affected_tonnesC",
    }
    range_sql = _range_sum_expressions(
        value_expression="affected_kgC",
        unit_stems=unit_stems,
        units=units,
        hide_ranges=hide_ranges,
        estimate_case_expression="estimate_case",
    )
    outer_select = _join_sql_parts([
        ", ".join(_quote_ident(name) for name in group_names),
        "carbon_origin",
        *range_sql,
    ])
    group_columns = [*group_names, "carbon_origin"]

    sql = f"""
        WITH direct_carbon AS (
            {' UNION ALL '.join(union_parts)}
        )
        SELECT {outer_select}
        FROM direct_carbon
        {_group_by_sql(group_columns)}
        {_order_by_sql(group_columns)};
    """

    rows = _fetch_dicts(conn, sql, union_params)
    incident_count = _count_distinct_incidents(
        conn,
        table=TABLE_STAGE1,
        where_parts=where_parts,
        params=params,
    )
    range_columns = _range_column_names(
        unit_stems=unit_stems,
        units=units,
        hide_ranges=hide_ranges,
    )
    plot_stem = unit_stems[units[0]]

    return ReportTable(
        stage="stage1",
        name="stage1_affected_carbon",
        title="Stage 1 directly affected carbon stock",
        rows=rows,
        columns=[*group_names, "carbon_origin", *range_columns],
        applied_groupings=_group_names(applied),
        skipped_groupings=skipped,
        incident_count=incident_count,
        plot_value_column=f"{plot_stem}_estimate",
        plot_lower_column=None if hide_ranges else f"{plot_stem}_lower",
        plot_upper_column=None if hide_ranges else f"{plot_stem}_upper",
        plot_value_label=(
            "Affected carbon (kgC)" if units[0] == "kg" else "Affected carbon (tonnes C)"
        ),
        plot_category_columns=[*group_names, "carbon_origin"],
        plot_series_column=None,
        notes=[
            "Stage 1 represents affected carbon stock before combustion completeness and species conversion.",
            "Total carbon overlaps with the biogenic and fossil rows; do not sum all three origins together.",
        ],
    )

def _build_replacement_table(
    conn: sqlite3.Connection,
    *,
    metadata: dict[str, Any],
    requested_groupings: Sequence[str],
    hide_ranges: bool,
    units: Sequence[str],
) -> ReportTable:
    _require_table(conn, TABLE_STAGE1)

    applied, skipped = _resolve_groupings_for_table(
        conn,
        stage="replacement",
        table=TABLE_STAGE1,
        requested_groupings=requested_groupings,
    )
    physical = _physical_groupings(applied)
    group_names = _group_names(physical)
    group_projection = _group_projection(physical)

    columns = _table_columns(conn, TABLE_STAGE1)
    where_parts = ["LOWER(COALESCE(src.emission_pathway, '')) = 'replacement'"]
    params: list[Any] = []

    input_where, input_params = _input_type_filter(columns, metadata)
    where_parts.extend(input_where)
    where_parts.extend(_calculation_status_clause(columns))
    params.extend(input_params)

    unit_stems = {
        "kg": "replacement_CO2_kg",
        "tonnes": "replacement_CO2_tonnes",
    }
    range_sql = _range_sum_expressions(
        value_expression="COALESCE(src.replacement_embodied_CO2_kg, 0.0)",
        unit_stems=unit_stems,
        units=units,
        hide_ranges=hide_ranges,
    )
    select_sql = _join_sql_parts([group_projection, *range_sql])

    sql = f"""
        SELECT {select_sql}
        FROM {_quote_ident(TABLE_STAGE1)} AS src
        WHERE {' AND '.join(where_parts)}
        {_group_by_sql(group_names)}
        {_order_by_sql(group_names)};
    """

    rows = _fetch_dicts(conn, sql, params)
    incident_count = _count_distinct_incidents(
        conn,
        table=TABLE_STAGE1,
        where_parts=where_parts,
        params=params,
    )
    range_columns = _range_column_names(
        unit_stems=unit_stems,
        units=units,
        hide_ranges=hide_ranges,
    )
    plot_stem = unit_stems[units[0]]

    return ReportTable(
        stage="replacement",
        name="replacement_CO2",
        title="Embodied CO2 emissions",
        rows=rows,
        columns=[*group_names, *range_columns],
        applied_groupings=_group_names(applied),
        skipped_groupings=skipped,
        incident_count=incident_count,
        plot_value_column=f"{plot_stem}_estimate",
        plot_lower_column=None if hide_ranges else f"{plot_stem}_lower",
        plot_upper_column=None if hide_ranges else f"{plot_stem}_upper",
        plot_value_label=(
            "Replacement CO2 emissions (kg)"
            if units[0] == "kg"
            else "Replacement CO2 emissions (tonnes)"
        ),
        plot_category_columns=group_names,
        plot_series_column=None,
        description=(
            "These are the estimated CO2 emissions required to produce "
            "replacements for fire-damaged items."
        ),
        notes=[
            "Replacement embodied CO2 is read from Stage 1 but reported as its own final-impact stream.",
            "It is not a combustion species and is not included in Stage 2 chemistry.",
        ],
    )

def _build_stage2_table(
    conn: sqlite3.Connection,
    *,
    metadata: dict[str, Any],
    requested_groupings: Sequence[str],
    hide_ranges: bool,
    units: Sequence[str],
) -> ReportTable:
    _require_table(conn, TABLE_STAGE2)
    _require_table(conn, TABLE_STAGE1)

    applied, skipped = _resolve_groupings_for_table(
        conn,
        stage="stage2",
        table=TABLE_STAGE2,
        requested_groupings=requested_groupings,
    )
    physical = _physical_groupings(
        applied,
        intrinsic_aliases={"species", "carbon_origin"},
    )
    group_names = _group_names(physical)
    group_projection = _group_projection(physical)

    columns = _table_columns(conn, TABLE_STAGE2)
    where_parts: list[str] = []
    params: list[Any] = []

    input_where, input_params = _input_type_filter(columns, metadata)
    where_parts.extend(input_where)
    where_parts.extend(_calculation_status_clause(columns))
    params.extend(input_params)
    where_sql = "WHERE " + " AND ".join(where_parts) if where_parts else ""

    unit_stems = {"kg": "emitted_kg", "tonnes": "emitted_tonnes"}
    range_sql = _range_sum_expressions(
        value_expression="COALESCE(src.emitted_kg, 0.0)",
        unit_stems=unit_stems,
        units=units,
        hide_ranges=hide_ranges,
    )
    select_sql = _join_sql_parts([
        group_projection,
        "src.emission_species",
        "src.carbon_origin",
        *range_sql,
    ])
    group_columns = [*group_names, "emission_species", "carbon_origin"]

    sql = f"""
        SELECT {select_sql}
        FROM {_quote_ident(TABLE_STAGE2)} AS src
        {where_sql}
        {_group_by_sql(group_columns)}
        {_order_by_sql(group_columns)};
    """

    rows = _fetch_dicts(conn, sql, params)

    # Display CO2 before CO while preserving the requested grouping order.
    species_order = {"CO2": 0, "CO": 1}
    rows.sort(
        key=lambda row: (
            tuple(str(row.get(name) or "") for name in group_names),
            species_order.get(str(row.get("emission_species") or ""), 99),
            str(row.get("carbon_origin") or ""),
        )
    )

    # Count the Stage 1 direct-pathway incidents before Stage 2's positive-value
    # filter and species-factor omissions.  This provides one stable stage-level
    # count instead of different counts for low/default/high output rows.
    stage1_columns = _table_columns(conn, TABLE_STAGE1)
    count_where = ["LOWER(COALESCE(src.emission_pathway, '')) = 'direct'"]
    count_params: list[Any] = []
    input_where, input_params = _input_type_filter(stage1_columns, metadata)
    count_where.extend(input_where)
    count_where.extend(_calculation_status_clause(stage1_columns))
    count_params.extend(input_params)
    incident_count = _count_distinct_incidents(
        conn,
        table=TABLE_STAGE1,
        where_parts=count_where,
        params=count_params,
    )

    range_columns = _range_column_names(
        unit_stems=unit_stems,
        units=units,
        hide_ranges=hide_ranges,
    )
    plot_stem = unit_stems[units[0]]

    return ReportTable(
        stage="stage2",
        name="stage2_direct_emissions",
        title="Stage 2 direct combustion emissions",
        rows=rows,
        columns=[
            *group_names,
            "emission_species",
            "carbon_origin",
            *range_columns,
        ],
        applied_groupings=_group_names(applied),
        skipped_groupings=skipped,
        incident_count=incident_count,
        plot_value_column=f"{plot_stem}_estimate",
        plot_lower_column=None if hide_ranges else f"{plot_stem}_lower",
        plot_upper_column=None if hide_ranges else f"{plot_stem}_upper",
        plot_value_label=(
            "Direct emitted mass (kg)" if units[0] == "kg" else "Direct emitted mass (tonnes)"
        ),
        plot_category_columns=[
            *group_names,
            "emission_species",
            "carbon_origin",
        ],
        plot_series_column=None,
        notes=[
            "Use carbon_origin='total' for overall species totals. Biogenic and fossil rows are partitions of the total and must not be added to it.",
            "Replacement embodied CO2 is deliberately excluded and reported by the separate replacement stage.",
        ],
    )

# =============================================================================
# TERMINAL TABLE OUTPUT
# =============================================================================

def _format_value(value: Any) -> str:
    """Format values compactly for terminal and Word tables."""
    if value is None:
        return ""
    if isinstance(value, bool):
        return "yes" if value else "no"
    if isinstance(value, int):
        return f"{value:,}"
    if isinstance(value, float):
        if not math.isfinite(value):
            return str(value)
        if abs(value) >= 1000:
            return f"{value:,.2f}"
        if abs(value) >= 1:
            return f"{value:,.3f}"
        return f"{value:,.6f}".rstrip("0").rstrip(".")
    return str(value)


def _print_metadata(result: ReportResult) -> None:
    metadata = result.metadata

    print("\nStored model build")
    print("------------------")
    if not metadata:
        print("No fire_model_metadata row was found.")
        return

    fields = [
        ("Model", metadata.get("model_name")),
        ("Version", metadata.get("model_version")),
        ("Input type", metadata.get("input_type")),
        ("Inventory snapshot", metadata.get("inventory_snapshot_id")),
        ("Completed", metadata.get("finished_utc") or metadata.get("created_at_utc")),
    ]

    for label, value in fields:
        if value is not None:
            print(f"{label + ':':<20} {value}")


def _print_report_table(table: ReportTable) -> None:
    # Two leading line breaks make adjacent report sections easier to separate.
    print("\n\n" + table.title)
    print("-" * len(table.title))

    if table.description:
        print(table.description)

    for summary_line in table.summary_lines:
        print(summary_line)

    if table.applied_groupings:
        print("Applied groupings: " + " > ".join(table.applied_groupings))
    else:
        print("Applied groupings: none")

    if table.skipped_groupings:
        print("Skipped groupings: " + ", ".join(table.skipped_groupings))

    if table.incident_count is not None:
        print(f"{table.incident_count_label} = {table.incident_count:,}")

    # Leave one blank line between the table context and its column header.
    print()

    if not table.rows:
        print("No rows were returned.")
        return

    display_rows = table.rows[:MAX_CONSOLE_ROWS]

    # Calculate readable column widths from the displayed rows only.
    widths: dict[str, int] = {}
    for column in table.columns:
        max_value_width = max(
            [len(_format_value(row.get(column))) for row in display_rows] or [0]
        )
        maximum_width = table.console_column_max_widths.get(column, 42)
        widths[column] = min(max(len(column), max_value_width), maximum_width)

    def shorten(text: str, width: int) -> str:
        if len(text) <= width:
            return text
        if width <= 3:
            return text[:width]
        return text[: width - 3] + "..."

    header = " | ".join(
        shorten(column, widths[column]).ljust(widths[column])
        for column in table.columns
    )
    divider = "-+-".join("-" * widths[column] for column in table.columns)

    print(header)
    print(divider)

    for row in display_rows:
        print(
            " | ".join(
                shorten(_format_value(row.get(column)), widths[column]).ljust(widths[column])
                for column in table.columns
            )
        )

    if len(table.rows) > len(display_rows):
        print(
            f"... terminal display limited to {len(display_rows):,} of "
            f"{len(table.rows):,} rows."
        )

def print_report(result: ReportResult) -> None:
    """Print the Stage 3 summaries to the terminal."""
    print("\nFire Emissions model report")
    print("===========================")
    print(f"Database: {result.db_path}")
    print("Stages: " + ", ".join(result.requested_stages))
    print(
        "Requested grouping hierarchy: "
        + (" > ".join(result.requested_groupings) if result.requested_groupings else "none")
    )
    print("Ranges: " + ("default only" if result.hide_ranges else "estimate, lower and upper"))
    print("Units: " + ", ".join(result.units))

    _print_metadata(result)

    for table in result.tables:
        _print_report_table(table)

# =============================================================================
# PLOT OUTPUT
# =============================================================================

def _plot_category_label(
    row: dict[str, Any],
    category_columns: Sequence[str],
) -> str:
    if not category_columns:
        return "Total"

    parts = []
    for column in category_columns:
        value = row.get(column)
        text = "(missing)" if value in (None, "") else str(value)
        parts.append(text)
    return " | ".join(parts)


def _estimate_case_order(values: Iterable[str]) -> list[str]:
    preferred = ["low", "default", "high"]
    unique = list(dict.fromkeys(str(value) for value in values))
    return [value for value in preferred if value in unique] + [
        value for value in unique if value not in preferred
    ]


def create_plots(
    result: ReportResult,
    *,
    output_dir: Optional[Path] = None,
) -> list[tuple[ReportTable, Any, Optional[Path]]]:
    """
    Create one straightforward summary plot for each plottable report table.

    Range-bearing tables plot their default estimate.  When lower and upper
    columns are visible, they are drawn as asymmetric error bars around that
    estimate.  The first requested unit is used for plots when the table
    includes both kilograms and tonnes.
    """
    try:
        import matplotlib.pyplot as plt
    except ImportError as exc:  # pragma: no cover - depends on local environment
        raise RuntimeError(
            "Plot output requires matplotlib. Install it in the active Python "
            "environment before using --plots or --write plots."
        ) from exc

    if output_dir is not None:
        output_dir.mkdir(parents=True, exist_ok=True)

    plotted: list[tuple[ReportTable, Any, Optional[Path]]] = []

    for table in result.tables:
        value_column = table.plot_value_column
        if value_column is None or not table.rows:
            continue

        series_column = table.plot_series_column
        category_columns = table.plot_category_columns

        values: dict[str, dict[str, float]] = defaultdict(lambda: defaultdict(float))
        lower_values: dict[str, dict[str, float]] = defaultdict(lambda: defaultdict(float))
        upper_values: dict[str, dict[str, float]] = defaultdict(lambda: defaultdict(float))

        for row in table.rows:
            raw_value = row.get(value_column)
            if raw_value is None:
                continue

            category = _plot_category_label(row, category_columns)
            series = (
                str(row.get(series_column) or "value")
                if series_column is not None
                else "value"
            )
            values[category][series] += float(raw_value)

            if table.plot_lower_column is not None:
                raw_lower = row.get(table.plot_lower_column)
                if raw_lower is not None:
                    lower_values[category][series] += float(raw_lower)

            if table.plot_upper_column is not None:
                raw_upper = row.get(table.plot_upper_column)
                if raw_upper is not None:
                    upper_values[category][series] += float(raw_upper)

        if not values:
            continue

        series_names = list(dict.fromkeys(
            series
            for category_values in values.values()
            for series in category_values.keys()
        ))

        def category_rank(item: tuple[str, dict[str, float]]) -> float:
            _, category_values = item
            return max(category_values.values(), default=0.0)

        ranked = sorted(values.items(), key=category_rank, reverse=True)
        ranked = ranked[:MAX_PLOT_CATEGORIES]

        categories = [category for category, _ in ranked]
        category_values = dict(ranked)

        use_horizontal = len(categories) > 8 or any(len(c) > 24 for c in categories)
        figure_height = max(5.0, min(18.0, 0.38 * len(categories) + 2.5))
        fig, ax = plt.subplots(figsize=(11, figure_height if use_horizontal else 6.5))

        n_series = max(1, len(series_names))
        bar_width = 0.8 / n_series
        positions = list(range(len(categories)))

        for series_index, series_name in enumerate(series_names):
            offset = (series_index - (n_series - 1) / 2.0) * bar_width
            plotted_values = [
                category_values[category].get(series_name, 0.0)
                for category in categories
            ]

            error_values = None
            if table.plot_lower_column is not None and table.plot_upper_column is not None:
                lower_errors = []
                upper_errors = []
                for category, estimate in zip(categories, plotted_values):
                    lower = lower_values[category].get(series_name, estimate)
                    upper = upper_values[category].get(series_name, estimate)
                    lower_errors.append(max(0.0, estimate - lower))
                    upper_errors.append(max(0.0, upper - estimate))
                error_values = [lower_errors, upper_errors]

            shifted = [position + offset for position in positions]
            if use_horizontal:
                ax.barh(
                    shifted,
                    plotted_values,
                    height=bar_width,
                    xerr=error_values,
                    capsize=3 if error_values is not None else 0,
                    label=series_name if n_series > 1 else None,
                )
            else:
                ax.bar(
                    shifted,
                    plotted_values,
                    width=bar_width,
                    yerr=error_values,
                    capsize=3 if error_values is not None else 0,
                    label=series_name if n_series > 1 else None,
                )

        title = table.title
        if len(values) > MAX_PLOT_CATEGORIES:
            title += f" (top {MAX_PLOT_CATEGORIES})"
        ax.set_title(title)

        if use_horizontal:
            ax.set_yticks(positions)
            ax.set_yticklabels(categories)
            ax.invert_yaxis()
            ax.set_xlabel(table.plot_value_label or value_column)
        else:
            ax.set_xticks(positions)
            ax.set_xticklabels(categories, rotation=45, ha="right")
            ax.set_ylabel(table.plot_value_label or value_column)

        if n_series > 1:
            ax.legend(title=series_column)

        fig.tight_layout()

        plot_path: Optional[Path] = None
        if output_dir is not None:
            plot_path = output_dir / f"{table.name}.png"
            fig.savefig(plot_path, dpi=180, bbox_inches="tight")

        plotted.append((table, fig, plot_path))

    return plotted

# =============================================================================
# WORD DOCUMENT OUTPUT
# =============================================================================

def _normalise_write_values(values: Optional[Sequence[str]]) -> set[str]:
    """
    Resolve the special --write behaviour.

    * option absent      -> no document
    * --write            -> tables + plots
    * --write all        -> tables + plots
    * --write tables     -> tables only
    * --write plots      -> plots only
    """
    if values is None:
        return set()
    if len(values) == 0 or "all" in values:
        return {"tables", "plots"}
    return set(values)


def write_report_document(
    result: ReportResult,
    *,
    write_modes: set[str],
    plots: Sequence[tuple[ReportTable, Any, Optional[Path]]],
) -> Path:
    """Write the requested tables and/or plots to one Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:  # pragma: no cover - depends on local environment
        raise RuntimeError(
            "Word output requires python-docx. Install it in the active Python "
            "environment before using --write."
        ) from exc

    output_root = result.db_path.parent / "reports"
    output_root.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    document_path = output_root / f"fire_model_report_{timestamp}.docx"

    doc = Document()
    doc.add_heading("Fire Emissions model report", level=0)

    doc.add_paragraph(f"Database: {result.db_path}")
    doc.add_paragraph("Stages: " + ", ".join(result.requested_stages))
    doc.add_paragraph(
        "Grouping hierarchy: "
        + (" > ".join(result.requested_groupings) if result.requested_groupings else "none")
    )
    doc.add_paragraph(
        "Ranges: " + ("default only" if result.hide_ranges else "estimate, lower and upper")
    )
    doc.add_paragraph("Units: " + ", ".join(result.units))

    if result.metadata:
        doc.add_heading("Stored model build", level=1)
        metadata_fields = [
            ("Model", result.metadata.get("model_name")),
            ("Version", result.metadata.get("model_version")),
            ("Input type", result.metadata.get("input_type")),
            ("Inventory snapshot", result.metadata.get("inventory_snapshot_id")),
            (
                "Completed",
                result.metadata.get("finished_utc")
                or result.metadata.get("created_at_utc"),
            ),
        ]
        for label, value in metadata_fields:
            if value is not None:
                doc.add_paragraph(f"{label}: {value}")

    if "tables" in write_modes:
        doc.add_heading("Summary tables", level=1)

        for report_table in result.tables:
            # Add a blank paragraph so consecutive report sections are clearly
            # separated even when the document style uses tight heading spacing.
            doc.add_paragraph("")
            doc.add_heading(report_table.title, level=2)

            if report_table.description:
                doc.add_paragraph(report_table.description)

            for summary_line in report_table.summary_lines:
                doc.add_paragraph(summary_line)

            applied_text = (
                " > ".join(report_table.applied_groupings)
                if report_table.applied_groupings
                else "none"
            )
            doc.add_paragraph(f"Applied groupings: {applied_text}")

            if report_table.skipped_groupings:
                doc.add_paragraph(
                    "Skipped groupings: "
                    + ", ".join(report_table.skipped_groupings)
                )

            if report_table.incident_count is not None:
                doc.add_paragraph(
                    f"{report_table.incident_count_label} = "
                    f"{report_table.incident_count:,}"
                )

            if not report_table.rows:
                doc.add_paragraph("No rows were returned.")
                continue

            rows_to_write = report_table.rows[:MAX_DOCUMENT_ROWS]
            # Separate the explanatory text from the table header itself.
            doc.add_paragraph("")
            table = doc.add_table(rows=1, cols=len(report_table.columns))
            table.style = "Table Grid"

            for index, column in enumerate(report_table.columns):
                table.rows[0].cells[index].text = column

            for row in rows_to_write:
                cells = table.add_row().cells
                for index, column in enumerate(report_table.columns):
                    cells[index].text = _format_value(row.get(column))

            if len(report_table.rows) > len(rows_to_write):
                doc.add_paragraph(
                    f"Document table limited to {len(rows_to_write):,} of "
                    f"{len(report_table.rows):,} rows."
                )

    if "plots" in write_modes:
        doc.add_heading("Summary plots", level=1)

        if not plots:
            doc.add_paragraph("No plottable summary tables were produced.")

        for report_table, _figure, plot_path in plots:
            doc.add_heading(report_table.title, level=2)
            if plot_path is None:
                doc.add_paragraph("Plot file was not available.")
                continue
            doc.add_picture(str(plot_path), width=Inches(6.5))

    doc.save(document_path)
    return document_path


# =============================================================================
# COMMAND-LINE INTERFACE
# =============================================================================

def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        prog="model_report",
        description=(
            "Summarise existing Fire Emissions model tables without rebuilding "
            "or overwriting the model."
        ),
    )

    parser.add_argument(
        "--profile",
        required=True,
        help="Profile name from config/local_paths.yaml, for example tom.",
    )

    parser.add_argument(
        "--db",
        required=True,
        help="Database handle from config/local_paths.yaml, normally fire_db.",
    )

    parser.add_argument(
        "--stage",
        nargs="+",
        choices=[*STAGE_ORDER, "all"],
        default=list(DEFAULT_STAGES),
        help=(
            "Logical report stages to create. Default: stage2 replacement. "
            "Use 'all' for inventory event stage1 replacement stage2."
        ),
    )

    parser.add_argument(
        "--group-by",
        nargs="+",
        choices=["none", *GROUPING_REGISTRY.keys()],
        default=["none"],
        help=(
            "Optional grouping hierarchy, applied left to right. Default: none. "
            "Incompatible groupings are skipped for the affected summary and reported."
        ),
    )

    parser.add_argument(
        "--units",
        nargs="+",
        choices=UNIT_ORDER,
        default=list(DEFAULT_UNITS),
        help=(
            "Mass-unit mode for range-bearing summary tables. Default: "
            "'default', which uses kgC for inventory room carbon stock and "
            "tonnes for all other mass summaries. Supply 'kg', 'tonnes', or "
            "'kg tonnes' to override the stage defaults everywhere."
        ),
    )

    parser.add_argument(
        "--hide-ranges",
        action="store_true",
        help=(
            "Show existing default estimate rows only. This hides low/high from "
            "the output but does not rerun or change the model."
        ),
    )

    parser.add_argument(
        "--plots",
        action="store_true",
        help="Create plots for every requested stage using its applied groupings.",
    )

    parser.add_argument(
        "--write",
        nargs="*",
        choices=["tables", "plots", "all"],
        default=None,
        help=(
            "Write a Word document. If supplied without values, writes tables "
            "and plots. Explicit values: tables, plots, or all."
        ),
    )

    return parser.parse_args(argv)


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)

    try:
        stages = _normalise_stages(args.stage)
        groupings = _normalise_groupings(args.group_by)
        units = _normalise_units(args.units)
        write_modes = _normalise_write_values(args.write)

        # --write plots (or --write with no following values) automatically
        # enables plot creation, as agreed for the reporting CLI.
        plots_required = args.plots or "plots" in write_modes

        config = load_local_paths_config(Path("config") / "local_paths.yaml")
        resolved = resolve_db_path(args.profile, args.db, config)
        db_path = Path(resolved.db_path)

        print("Resolved reporting request:")
        print(f"  DB HANDLE:  {args.db}")
        print(f"  DB:         {db_path}")
        print(f"  STAGES:     {'; '.join(stages)}")
        print(
            "  GROUPINGS:  "
            + (" > ".join(groupings) if groupings else "none")
        )
        print(f"  RANGES:     {'default only' if args.hide_ranges else 'estimate; lower; upper'}")
        print(f"  UNITS:      {'; '.join(units)}")
        print(f"  PLOTS:      {'yes' if plots_required else 'no'}")
        print(
            "  WRITE:      "
            + ("; ".join(sorted(write_modes)) if write_modes else "none")
        )

        if not db_path.exists():
            print("\nERROR: Database file does not exist at the resolved path.")
            return 2

        result = build_model_report(
            db_path,
            stages=stages,
            group_by=groupings or ("none",),
            hide_ranges=args.hide_ranges,
            units=units,
        )

        print_report(result)

        plot_results: list[tuple[ReportTable, Any, Optional[Path]]] = []
        document_path: Optional[Path] = None

        if plots_required:
            plot_output_dir: Optional[Path] = None
            if "plots" in write_modes:
                timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
                plot_output_dir = (
                    db_path.parent
                    / "reports"
                    / "model_report_assets"
                    / timestamp
                )

            plot_results = create_plots(
                result,
                output_dir=plot_output_dir,
            )

        if write_modes:
            document_path = write_report_document(
                result,
                write_modes=write_modes,
                plots=plot_results,
            )
            print(f"\nReport document written to: {document_path}")

        # Show interactive plot windows only when --plots was explicitly used.
        # --write plots creates/saves plots silently unless --plots is also present.
        if args.plots and plot_results:
            import matplotlib.pyplot as plt

            plt.show()

        print("\nReporting complete.")
        print(f"  Summary tables produced: {len(result.tables)}")
        print(f"  Plots produced:          {len(plot_results)}")
        written_document_count = 1 if document_path is not None else 0
        print(f"  Documents written:       {written_document_count}")
        return 0

    except Exception as exc:
        print("\nERROR:", exc)
        return 3


if __name__ == "__main__":
    raise SystemExit(main())
